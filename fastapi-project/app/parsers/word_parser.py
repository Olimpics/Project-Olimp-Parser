from typing import Dict, Any, List
import docx
import re

async def parse_disciplines(file_path: str, limit: int = 5) -> List[Dict[str, Any]]:
    try:
        doc = docx.Document(file_path)
        
        paragraphs = doc.paragraphs[:min(100, len(doc.paragraphs))]
        text = "\n".join([paragraph.text for paragraph in paragraphs])
        
        disciplines = []
        
        discipline_blocks = re.split(r'\n\s*\n', text)
        
        for block in discipline_blocks:
            if re.search(r'назва дисципліни|код дисципліни', block, re.IGNORECASE):
                name_match = re.search(r'назва(?:\s+дисципліни)?[:\s]+([^\n]+)', block, re.IGNORECASE)
                name = name_match.group(1).strip() if name_match else ""
                
                code_match = re.search(r'код(?:\s+дисципліни)?[:\s]+([^\n]+)', block, re.IGNORECASE)
                code = code_match.group(1).strip() if code_match else ""
                
                faculty_match = re.search(r'факультет[:\s]+([^\n]+)', block, re.IGNORECASE)
                faculty = faculty_match.group(1).strip() if faculty_match else ""
                
                min_count_match = re.search(r'мін(?:імальна)?(?:\s+кількість)?[:\s]+(\d+)', block, re.IGNORECASE)
                min_count = int(min_count_match.group(1)) if min_count_match else 0
                
                max_count_match = re.search(r'макс(?:имальна)?(?:\s+кількість)?[:\s]+(\d+)', block, re.IGNORECASE)
                max_count = int(max_count_match.group(1)) if max_count_match else 0
                
                details = {
                    "departmentId": 0,
                    "teacher": "",
                    "recomend": "",
                    "prerequisites": "",
                    "language": "",
                    "determination": "",
                    "whyInterestingDetermination": "",
                    "resultEducation": "",
                    "usingIrl": "",
                    "additionaLiterature": "",
                    "typesOfTraining": "",
                    "typeOfControll": ""
                }
                
                teacher_match = re.search(r'викладач[:\s]+([^\n]+)', block, re.IGNORECASE)
                if teacher_match:
                    details["teacher"] = teacher_match.group(1).strip()
                
                discipline = {
                    "nameAddDisciplines": name,
                    "codeAddDisciplines": code,
                    "faculty": faculty,
                    "minCountPeople": min_count,
                    "maxCountPeople": max_count,
                    "minCourse": 0,
                    "maxCourse": 0,
                    "addSemestr": "",
                    "degreeLevel": "",
                    "details": details,
                    "idAddDisciplines": 0
                }
                disciplines.append(discipline)
                
                if len(disciplines) >= limit:
                    break
        
        return disciplines[:limit]
    
    except Exception as e:
        print(f"Помилка при парсингу Word файлу дисциплін: {str(e)}")
        raise ValueError(f"Не вдалося розібрати Word файл дисциплін: {str(e)}")

async def parse_educational_programs(file_path: str) -> Dict[str, Any]:
    try:
        doc = docx.Document(file_path)
        
        program_name = ""
        program_name_en = ""
        degree = ""
        accreditation_type = ""
        credits = 0
        speciality = ""
        
        semester_counts = {3: 0, 4: 0, 5: 0, 6: 0, 7: 0, 8: 0}
        
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                    
                header_cell = row.cells[0].text.strip().lower()
                value_cell = row.cells[1].text.strip()
                
                if ("офіційна назва освітньої програми" in header_cell or 
                    "повна назва освітньої програми" in header_cell) and "англійськ" not in header_cell:
                    program_name = value_cell
                    print(f"Found Ukrainian program name in table: {program_name}")
                
                elif ("офіційна назва освітньої програми" in header_cell or 
                      "назва освітньої програми" in header_cell) and "англійськ" in header_cell:
                    program_name_en = value_cell
                    print(f"Found English program name in table: {program_name_en}")
                
                elif re.search(r'ступінь.*освіти|кваліфікація', header_cell):
                    degree_match = re.search(r'бакалавр|магістр|доктор філософії', value_cell, re.IGNORECASE)
                    if degree_match:
                        degree = degree_match.group(0).lower()
                
                elif re.search(r'спеціальність', header_cell):
                    speciality_match = re.search(r'\d+\s+\w+', value_cell)
                    if speciality_match:
                        speciality = speciality_match.group(0)
                
                elif re.search(r'акредитаці[яї]|наявність', header_cell):
                    accreditation_type = value_cell
                
                elif re.search(r'обсяг|диплом', header_cell) and 'кредит' in value_cell.lower():
                    credits_match = re.search(r'(\d+)\s*кредит', value_cell)
                    if credits_match:
                        credits = int(credits_match.group(1))
        
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        if not program_name:
            ukr_patterns = [
                r'Освітньо-професійна програма «([^»]+)»',
                r'Освітня програма: «([^»]+)»',
                r'Офіційна назва освітньої програми[:\s]+([^\n]+)'
            ]
            
            for pattern in ukr_patterns:
                ukr_match = re.search(pattern, text, re.IGNORECASE)
                if ukr_match:
                    if "Освітньо-професійна програма" in ukr_match.group(0):
                        program_name = f"Освітньо-професійна програма «{ukr_match.group(1).strip()}»"
                    elif "Освітня програма" in ukr_match.group(0):
                        program_name = f"Освітня програма: «{ukr_match.group(1).strip()}»"
                    else:
                        program_name = ukr_match.group(1).strip()
                    print(f"Found Ukrainian program name in text: {program_name}")
                    break
        
        if not degree:
            degree_match = re.search(r'ступінь[:\s]+([^\n]+)', text, re.IGNORECASE)
            if degree_match:
                degree = degree_match.group(1).strip().lower()
        
        if not speciality:
            speciality_match = re.search(r'спеціальність[:\s]+(\d+[^\n,;]+)', text, re.IGNORECASE)
            if speciality_match:
                speciality = speciality_match.group(1).strip()
        
        if not credits:
            credits_match = re.search(r'(\d+)\s*кредит', text)
            if credits_match:
                credits = int(credits_match.group(1))
        
        final_program_name = program_name if program_name else "Освітня програма"
        
        for table in doc.tables:
            is_electives_table = False
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(re.search(r'Вибіркові компоненти|ВК \d+', cell) for cell in cells):
                    is_electives_table = True
                    break
            
            if not is_electives_table:
                continue
                
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                
                if len(cells) < 3:
                    continue
                
                if cells[0].startswith("ВК"):
                    try:
                        semester = 0
                        for idx in range(len(cells) - 1, 0, -1):
                            if re.match(r'^\d+$', cells[idx]):
                                semester = int(cells[idx])
                                break
                        
                        if semester in semester_counts:
                            semester_counts[semester] += 1
                        
                    except (ValueError, IndexError):
                        pass
        
        print(f"Electives per semester: {semester_counts}")
        
        educational_program = {
            "idEducationalProgram": 0,
            "nameEducationalProgram": final_program_name,
            "countAddSemestr3": semester_counts[3],
            "countAddSemestr4": semester_counts[4],
            "countAddSemestr5": semester_counts[5],
            "countAddSemestr6": semester_counts[6],
            "countAddSemestr7": semester_counts[7],
            "countAddSemestr8": semester_counts[8],
            "degree": degree,
            "speciality": speciality,
            "accreditation": credits,
            "accreditationType": accreditation_type,
            "studentsAmount": 0,
            "studentsCount": 0,
            "disciplinesCount": 0
        }
        
        main_disciplines = []
        
        target_tables = []
        
        for table_idx, table in enumerate(doc.tables):
            found_cycle = False
            found_electives = False
            
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                
                if any("І Цикл загальної підготовки" in cell for cell in cells):
                    found_cycle = True
                
                if found_cycle and any("Вибіркові компоненти" in cell for cell in cells):
                    found_electives = True
            
            if found_cycle:
                target_tables.append(table_idx)
        
        for table_idx in target_tables:
            table = doc.tables[table_idx]
            
            is_collecting = False
            header_row = None
            
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                if any("І Цикл загальної підготовки" in cell for cell in cells):
                    is_collecting = True
                    if row_idx + 1 < len(table.rows):
                        header_cells = [cell.text.strip() for cell in table.rows[row_idx + 1].cells]
                        if any(re.search(r'форма.*контролю|семестр|кредити', cell, re.IGNORECASE) for cell in header_cells):
                            header_row = header_cells
                    continue
                
                if is_collecting and any(re.search(r'Вибіркові компоненти', cell, re.IGNORECASE) for cell in cells):
                    is_collecting = False
                    break
                
                if not is_collecting or len(cells) < 3:
                    continue
                
                if not cells[0].startswith("ОК"):
                    continue
                
                code = cells[0].strip()
                
                name_idx = 1
                form_control_idx = None
                credit_idx = None
                semester_idx = None
                
                if header_row:
                    for i, header in enumerate(header_row):
                        if re.search(r'форма.*контролю', header, re.IGNORECASE):
                            form_control_idx = i
                        elif re.search(r'кредити|ЄКТС', header, re.IGNORECASE):
                            credit_idx = i
                        elif re.search(r'семестр', header, re.IGNORECASE):
                            semester_idx = i
                
                if form_control_idx is None:
                    for i in range(2, min(5, len(cells))):
                        if re.search(r'екзамен|залік|іспит', cells[i], re.IGNORECASE):
                            form_control_idx = i
                            break
                
                if credit_idx is None:
                    for i in range(1, min(5, len(cells))):
                        if re.match(r'^\d+(?:[.,]\d+)?$', cells[i]):
                            credit_idx = i
                            break
                
                if semester_idx is None:
                    for i in range(len(cells) - 1, 0, -1):
                        if re.match(r'^\d+$', cells[i]) or re.match(r'^\d+(?:,\d+)+$', cells[i]):
                            semester_idx = i
                            break
                
                name_discipline = cells[name_idx] if name_idx < len(cells) else ""
                
                form_control = ""
                if form_control_idx is not None and form_control_idx < len(cells):
                    form_control = cells[form_control_idx]
                elif len(cells) > 2:
                    control_match = re.search(r'(екзамен|залік|іспит|диф\.\s*залік)', name_discipline, re.IGNORECASE)
                    if control_match:
                        form_control = control_match.group(1)
                        name_discipline = re.sub(r'\s*[(]?екзамен|залік|іспит|диф\.\s*залік[)]?\s*', '', name_discipline, flags=re.IGNORECASE)
                
                loans = 0
                if credit_idx is not None and credit_idx < len(cells):
                    loans_text = cells[credit_idx].replace(',', '.')
                    try:
                        loans = float(loans_text)
                    except ValueError:
                        loans = 0
                
                semestr = 0
                if semester_idx is not None and semester_idx < len(cells):
                    semestr_text = cells[semester_idx]
                    try:
                        semestr = int(semestr_text)
                    except ValueError:
                        if ',' in semestr_text:
                            try:
                                semestr = int(semestr_text.split(',')[0])
                            except ValueError:
                                semestr = 0
                
                discipline = {
                    "idBindMainDisciplines": 0,
                    "codeMainDisciplines": code,
                    "disciplineName": name_discipline,
                    "loans": loans,
                    "formControll": form_control,
                    "semestr": semestr,
                    "educationalProgramName": final_program_name
                }
                main_disciplines.append(discipline)
        
        if not main_disciplines:
            match = re.search(r'І Цикл загальної підготовки(.*?)Вибіркові компоненти', text, re.DOTALL | re.IGNORECASE)
            if match:
                disciplines_text = match.group(1)
                
                discipline_matches = re.finditer(
                    r'(ОК\s*\d+(?:\.\d+)?)\s+([^\n]+?)\s+(\d+(?:[.,]\d+)?)\s+(екзамен|залік|іспит|диф\.\s*залік)\s+(\d+(?:,\d+)?)',
                    disciplines_text,
                    re.IGNORECASE
                )
                
                for match in discipline_matches:
                    code = match.group(1).strip()
                    name_discipline = match.group(2).strip()
                    loans_text = match.group(3).replace(',', '.')
                    form_control = match.group(4).strip()
                    semestr_text = match.group(5)
                    
                    try:
                        loans = float(loans_text)
                    except ValueError:
                        loans = 0
                        
                    try:
                        if ',' in semestr_text:
                            semestr = int(semestr_text.split(',')[0])
                        else:
                            semestr = int(semestr_text)
                    except ValueError:
                        semestr = 0
                    
                    discipline = {
                        "idBindMainDisciplines": 0,
                        "codeMainDisciplines": code,
                        "disciplineName": name_discipline,
                        "loans": loans,
                        "formControll": form_control,
                        "semestr": semestr,
                        "educationalProgramName": final_program_name
                    }
                    main_disciplines.append(discipline)
        
        return {
            "educationalProgram": educational_program,
            "mainDisciplines": main_disciplines
        }
    
    except Exception as e:
        print(f"Помилка при парсингу Word файлу освітньої програми: {str(e)}")
        raise ValueError(f"Не вдалося розібрати Word файл освітньої програми: {str(e)}")